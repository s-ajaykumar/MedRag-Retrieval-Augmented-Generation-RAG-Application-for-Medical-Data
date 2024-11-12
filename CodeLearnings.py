#pip install faiss-cpu langchain_community uuid
from langchain_community.vectorstores import FAISS
from uuid import uuid4
from io import BytesIO

#Convert dictionary into Huggingface Dataset
from datasets import Dataset

def convert_analgesicDoc_into_dict():

    #To extract the contents of word document
    from docx import Document

    # Load the .docx file
    doc = Document('D:\Ajay\DeepLearning\GenAI\MedLLM\ETG-20240911T071246Z-001\ETG\Analgesic fix.docx')

    #Extract the paragraphs in the word document and store it in a single python string
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text.strip() + ' '

    #Extracting the title of the document
    title = doc.paragraphs[0].text

    #Store the document title and it's content in a dictionary
    dic = {'title' : [title], 'text' : [text]}

    #Create a HFDataset object for our dictionary
    dataset = Dataset.from_dict(dic)
    return dataset
dataset = convert_analgesicDoc_into_dict()


#To treat binary string as file object:
'''
#Read as binary string
with open("/Users/outbell/Ajay/DeepLearning/GenAI/MedLLM/VSCode/server_test_data/Analgesic fix copy.docx", 'rb') as f:
    file = f.read()

#Read as file object
file_object = BytesIO(file) '''



#Load the existing FAISS vector store
'''db = FAISS.load_local(
    "faiss_index", embedder, allow_dangerous_deserialization=True
)

#Index the documents and store in the vector store
uuids = [str(uuid4()) for _ in range(len(docs))]
db.add_documents(documents=docs, ids=uuids)

#Replace the new faiss vector store with the existing vectorstore in local
db.save_local("faiss_index")'''