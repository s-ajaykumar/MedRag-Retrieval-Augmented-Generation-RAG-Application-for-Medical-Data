#pip install pydantic langserve fastapi uvicorn sse_starlette
import base64

from fastapi import FastAPI
from langchain_core.runnables import RunnableLambda
from pydantic import Field
from langchain_core.document_loaders import BaseLoader
from typing import AsyncIterator, Iterator
from langchain_core.documents import Document
import os

from langserve import CustomUserType, add_routes


class CustomDocumentLoader(BaseLoader):
    """An example document loader that reads a file line by line."""

    def __init__(self, file_path: str) -> None:
        """Initialize the loader with a file path.

        Args:
            file_path: The path to the file to load.
        """
        self.file_path = file_path

    def lazy_load(self) -> Iterator[Document]:  # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """

        def split(title, subtitle, words):
          from langchain_core.documents import Document
          if subtitle is not None:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {content}"
                    )
          else:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {content}"
                    )

        def check_word_count(title, subtitle, passage):
          from langchain_core.documents import Document
          words = passage.split(' ')
          if len(words) > 90:
            yield from split(title, subtitle, words)
            return
          if subtitle is None:
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {passage}"
                    )
          else:
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {passage}"
                    )

        from docx import Document
        doc = Document(self.file_path)
        title = doc.paragraphs[0].text.strip()
        passage = ''

        for paragraph in doc.paragraphs:
          text = paragraph.text.strip()
          if text != '':
            count_of_spaces = 0
            for i in text:
              if i == ' ':
                count_of_spaces += 1
            if  count_of_spaces <= 5:
              if passage != '' and subtitle is None:
                yield from check_word_count(title, None, passage)
              elif passage != '' and subtitle is not None:
                yield from check_word_count(title, subtitle, passage)
              passage = ''
              subtitle = text
              continue
            passage += text + ' '


app = FastAPI(
    title="LangChain Server",
    version="1.0",
    description="Spin up a simple api server using Langchain's Runnable interfaces",
)



class FileProcessingRequest(CustomUserType):
    """Request including a base64 encoded file."""

    # The extra field is used to specify a widget for the playground UI.
    file: list = Field(...)


def _process_file(request: FileProcessingRequest) -> str:
    """Extract the text from the first page of the PDF."""

    # Count files in the directory
    def count_num_files(directory):

        file_count = len([file for file in os.listdir(directory) if os.path.isfile(os.path.join(directory, file))]) - 1
        return file_count

    file_count = count_num_files("/Users/outbell/Ajay/DeepLearning/GenAI/MedLLM/VSCode/server_test_data") #Pass the directory of the documents stored

    file_paths = []

    for req in request.file:
        content = base64.b64decode(req.encode("utf-8"))
        file_count += 1
        with open(f"/Users/outbell/Ajay/DeepLearning/GenAI/MedLLM/VSCode/server_test_data/Analgesic fix copy {file_count}.docx", "wb") as f:
            f.write(content)
        file_paths.append(f"/Users/outbell/Ajay/DeepLearning/GenAI/MedLLM/VSCode/server_test_data/Analgesic fix copy {file_count}.docx")
        
    n = 0
    for file_path in file_paths: 
        loader = CustomDocumentLoader(file_path)
        documents = loader.load()
        n += 1
    return f"{n} documents added. The total documents in the directory is {file_count}"


add_routes(
    app,
    RunnableLambda(_process_file).with_types(input_type=FileProcessingRequest),
    config_keys=["configurable"],
    path="/word",
)


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="localhost", port=8000)