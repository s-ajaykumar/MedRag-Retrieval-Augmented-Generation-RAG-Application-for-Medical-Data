from docx import Document
from typing import Iterator
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document

class CustomDocumentLoader(BaseLoader):
    """An example document loader that reads a file line by line."""

    def __init__(self, file_path: str) -> None:
        """Initialize the loader with a file path.

        Args:
            file_path: The path to the file to load.
        """
        self.file_path = file_path

    def lazy_load(self) -> Iterator[Document]:  # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """

        def split(title, subtitle, words):
          from langchain_core.documents import Document
          if subtitle is not None:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {content}"
                    )
          else:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {content}"
                    )

        def check_word_count(title, subtitle, passage):
          from langchain_core.documents import Document
          words = passage.split(' ')
          if len(words) > 90:
            yield from split(title, subtitle, words)
            return
          if subtitle is None:
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {passage}"
                    )
          else:
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {passage}"
                    )

        from docx import Document
        doc = Document(self.file_path)
        title = doc.paragraphs[0].text.strip()
        passage = ''

        for paragraph in doc.paragraphs:
          text = paragraph.text.strip()
          if text != '':
            count_of_spaces = 0
            for i in text:
              if i == ' ':
                count_of_spaces += 1
            if  count_of_spaces <= 5:
              if passage != '' and subtitle is None:
                yield from check_word_count(title, None, passage)
              elif passage != '' and subtitle is not None:
                yield from check_word_count(title, subtitle, passage)
              passage = ''
              subtitle = text
              continue
            passage += text + ' '