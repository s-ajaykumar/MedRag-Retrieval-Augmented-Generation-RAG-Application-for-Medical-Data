#pip install pydantic langserve fastapi uvicorn sse_starlette
import base64
from io import BytesIO

from fastapi import FastAPI
from langchain_core.runnables import RunnableLambda
from pydantic import Field
from langchain_core.document_loaders import BaseLoader
from typing import AsyncIterator, Iterator
from docx import Document
from azure.storage.blob import BlobServiceClient

from langserve import CustomUserType, add_routes

import logging
import config

import azure.functions as func
import azurefunctions.extensions.bindings.blob as blob


class CustomDocumentLoader(BaseLoader):
    """An example document loader that reads a file line by line."""

    def __init__(self, file_path: str) -> None:
        """Initialize the loader with a file path.

        Args:
            file_path: The path to the file to load.
        """
        self.file_path = file_path

    def lazy_load(self) -> Iterator[Document]:  # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """

        def split(title, subtitle, words):
          from langchain_core.documents import Document
          if subtitle is not None:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {content}"
                    )
          else:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {content}"
                    )

        def check_word_count(title, subtitle, passage):
          from langchain_core.documents import Document
          words = passage.split(' ')
          if len(words) > 90:
            yield from split(title, subtitle, words)
            return
          if subtitle is None:
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {passage}"
                    )
          else:
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {passage}"
                    )

        from docx import Document
        doc = Document(self.file_path)
        title = doc.paragraphs[0].text.strip()
        passage = ''

        for paragraph in doc.paragraphs:
          text = paragraph.text.strip()
          if text != '':
            count_of_spaces = 0
            for i in text:
              if i == ' ':
                count_of_spaces += 1
            if  count_of_spaces <= 5:
              if passage != '' and subtitle is None:
                yield from check_word_count(title, None, passage)
              elif passage != '' and subtitle is not None:
                yield from check_word_count(title, subtitle, passage)
              passage = ''
              subtitle = text
              continue
            passage += text + ' '


app = FastAPI(
    title="LangChain Server",
    version="1.0",
    description="Spin up a simple api server using Langchain's Runnable interfaces",
)


class FileProcessingRequest(CustomUserType):
    """Request including a base64 encoded file."""

    # The extra field is used to specify a widget for the playground UI.
    file: list = Field(...)


def _process_file(request: FileProcessingRequest) -> str:
    """Extract the text from the first page of the PDF."""


    #Authentication
    def get_blob_service_client_account_key():
    
        account_url = config.blob_url
        shared_access_key = config.blob_key
        credential = shared_access_key

        # Create the BlobServiceClient object
        blob_service_client = BlobServiceClient(account_url, credential=credential)

        return blob_service_client
    
    blob_service_client = get_blob_service_client_account_key()
    

    def count_num_blobs(blob_service_client: BlobServiceClient, container_name):
        l = 0
        container_client = blob_service_client.get_container_client(container=container_name)
        blob_list = container_client.list_blobs()

        for blob in blob_list:
            l += 1
        return l
      
    count_num_blobs = count_num_blobs(blob_service_client, "rawdocstore")


    def upload_blob_data(blob_service_client: BlobServiceClient, container_name: str, binary_data, count_num_blobs):
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=f"doc {count_num_blobs}.docx")

        # Upload the blob data - default blob type is BlockBlob
        blob_client.upload_blob(binary_data, blob_type="BlockBlob")


    binary_strings = []

    for req in request.file:
        content = base64.b64decode(req.encode("utf-8"))
        count_num_blobs += 1
        upload_blob_data(blob_service_client, "rawdocstore", content, count_num_blobs)
        binary_strings.append(content)
        
  
    for i in binary_strings: 
        file_object = BytesIO(i)
        loader = CustomDocumentLoader(file_object)
        docs = loader.load()
        
        
           


    app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)
    @app.blob_trigger(
        arg_name="client", path="rawdocstore/doc 3.docx", connection="AzureWebJobsStorage"
    )
    def blob_trigger(client: blob.BlobClient):
        logging.info(
            f"Python blob trigger function processed blob \n"
            f"Properties: {client.get_blob_properties()}\n"
            f"Blob content head: {client.download_blob().read(size=1)}"
        )
    


    return f"{len(binary_strings)} documents added. The total documents in the directory is {count_num_blobs}"


add_routes(
    app,
    RunnableLambda(_process_file).with_types(input_type=FileProcessingRequest),
    config_keys=["configurable"],
    path="/word",
)


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="localhost", port=8000)