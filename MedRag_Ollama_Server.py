#pip install langchain_ollama 
#Import libraries
from langchain_ollama import OllamaEmbeddings
from langchain_ollama import ChatOllama
#from langchain_core.messages import AIMessage
from typing import Iterator
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain import hub
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langserve import CustomUserType, add_routes
from fastapi import FastAPI
from pydantic import Field
from langchain_core.document_loaders import BaseLoader
from typing import Iterator
from docx import Document
from azure.storage.blob import BlobServiceClient
import base64
from io import BytesIO
from uuid import uuid4
from langchain_core.runnables import RunnableLambda
import json
import config


class blob_authentication:
   #Authentication
    def get_blob_service_client_account_key():
      account_url = config.blob_url
      shared_access_key = config.blob_key
      credential = shared_access_key

      # Create the BlobServiceClient object
      blob_service_client = BlobServiceClient(account_url, credential=credential)
      return blob_service_client
      
    blob_service_client = get_blob_service_client_account_key()

blob_key = blob_authentication()


class CustomDocumentLoader(BaseLoader):
    """An example document loader that reads a file line by line."""

    def __init__(self, file_path: str) -> None:
        """Initialize the loader with a file path.

        Args:
            file_path: The path to the file to load.
        """
        self.file_path = file_path

    def lazy_load(self) -> Iterator[Document]:  # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """

        def split(title, subtitle, words):
          from langchain_core.documents import Document
          if subtitle is not None:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {content}"
                    )
          else:
            for i in range(0, len(words), 90):
              content = ' '.join(words[i:i+90])
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {content}"
                    )

        def check_word_count(title, subtitle, passage):
          from langchain_core.documents import Document
          words = passage.split(' ')
          if len(words) > 90:
            yield from split(title, subtitle, words)
            return
          if subtitle is None:
              yield Document(
                        page_content = f"title: {title}, subtitle: {title}, passage: {passage}"
                    )
          else:
              yield Document(
                        page_content = f"title: {title}, subtitle: {subtitle}, passage: {passage}"
                    )

        from docx import Document
        doc = Document(self.file_path)
        title = doc.paragraphs[0].text.strip()
        passage = ''

        for paragraph in doc.paragraphs:
          text = paragraph.text.strip()
          if text != '':
            count_of_spaces = 0
            for i in text:
              if i == ' ':
                count_of_spaces += 1
            if  count_of_spaces <= 5:
              if passage != '' and subtitle is None:
                yield from check_word_count(title, None, passage)
              elif passage != '' and subtitle is not None:
                yield from check_word_count(title, subtitle, passage)
              passage = ''
              subtitle = text
              continue
            passage += text + ' '

class FileProcessingRequest(CustomUserType):
    """Request including a base64 encoded file."""

    # The extra field is used to specify a widget for the playground UI.
    file: list = Field(...)
    file_name: list = Field(...)


class rag:

  def __init__(self, blob_key):
      
      self.blob_key = blob_key
     #Load the embedder and llm
      self.embedder = OllamaEmbeddings(model = "mxbai-embed-large")
      self.llm = ChatOllama(model = "llama3.2")
      #Load the vector store
      #Download from azure blob storage and load it
      self.vector_store = Chroma(embedding_function = self.embedder, persist_directory = "/Users/outbell/Ajay/DeepLearning/GenAI/MedLLM/New_chromaDB")
      #initialize the retriever
      self.retriever = self.vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 10})
      #Prompt structure
      self.prompt = hub.pull("rlm/rag-prompt")


  #Concatenate retrieved passages
  def format_docs(self, docs):
          return "\n\n".join(doc.page_content for doc in docs)
     
  def upload_blob_data(self, container_name: str, file_object, file_name_string):
      container_client = self.blob_key.blob_service_client.get_container_client(container=container_name)
      # Upload the blob data - default blob type is BlockBlob
      container_client.upload_blob(name = file_name_string, data = file_object, overwrite = True)

  def _process_file(self, request: FileProcessingRequest) -> str:
      
      file_name_strings = []

      for name in request.file_name:
          file_name_bytes_data = base64.b64decode(name.encode("utf-8"))
          file_name_strings.append(file_name_bytes_data.decode('utf-8'))

      file_num = 0
      docs = []

      for req in request.file:
          content = base64.b64decode(req.encode("utf-8"))
          file_object = BytesIO(content)
          self.upload_blob_data("rawdocstore", file_object, file_name_strings[file_num])
          file_num += 1
          loader = CustomDocumentLoader(file_object)
          docs.extend(doc for doc in loader.load())
      uuids = [str(uuid4()) for _ in range(len(docs))]
      self.vector_store.add_documents(documents = docs, ids = uuids)
      #Reinitialize retriever
      self.retriever = self.vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 10})
      return "File/Files uploaded and vector_store updated"

#Create object for the class
rag_obj = rag(blob_key)
  
#Chain the components
rag_chain = (
    {"context": rag_obj.retriever | rag_obj.format_docs, "question": RunnablePassthrough()}
    | rag_obj.prompt
    | rag_obj.llm
    | StrOutputParser()
)

app = FastAPI(
    title="LangChain Server",
    version="1.0",
    description="Spin up a simple api server using Langchain's Runnable interfaces",
)
# Adds routes to the app for using the chain under:
# /invoke
# /batch
# /stream
add_routes(app, rag_chain, path = "/stream", enable_feedback_endpoint=True)
add_routes(app, 
           RunnableLambda(rag_obj._process_file).with_types(input_type=FileProcessingRequest),
           config_keys=["configurable"],
           path="/file_upload",)

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="localhost", port=8000)



